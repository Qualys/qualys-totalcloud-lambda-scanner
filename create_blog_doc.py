#!/usr/bin/env python3
"""
Generate a professionally formatted Word document for the Qualys Lambda Scanner blog post.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line after a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_image_placeholder(doc, caption):
    """Create a styled image placeholder box."""
    # Add placeholder paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[IMAGE PLACEHOLDER]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Add caption
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(10)
    caption_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacing

def main():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('', level=0)
    title_run = title.add_run('Securing AWS Lambda at Scale')
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('Continuous Vulnerability Management for Serverless Workloads')
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(102, 102, 102)
    sub_run.italic = True

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run('How organizations can extend their vulnerability management programs to cover the growing serverless attack surface')
    tag_run.font.size = Pt(12)
    tag_run.font.color.rgb = RGBColor(80, 80, 80)

    add_horizontal_line(tagline)
    doc.add_paragraph()

    # Hero image placeholder
    create_image_placeholder(doc, "Hero banner: AWS Lambda integration with Qualys platform - modern cloud security visual with professional design")

    # Section 1: The Serverless Shift
    doc.add_heading('The Serverless Shift and Its Security Implications', level=1)

    doc.add_paragraph(
        "The adoption of serverless computing continues to accelerate across enterprises of all sizes. "
        "AWS Lambda, the leading serverless platform, has fundamentally changed how organizations build "
        "and deploy applications. Development teams appreciate the reduced operational overhead, automatic "
        "scaling, and pay-per-execution pricing model. For many workloads, serverless represents a "
        "genuinely better way to build software."
    )

    doc.add_paragraph(
        "However, this architectural shift introduces new considerations for security teams. Traditional "
        "vulnerability management programs are built around persistent infrastructure—servers that exist "
        "long enough to be scanned, patched, and verified. Lambda functions operate differently. They may "
        "execute for milliseconds, scale to thousands of concurrent instances, and be redeployed multiple "
        "times per day. The tools and processes designed for traditional infrastructure don't translate "
        "directly to this model."
    )

    doc.add_paragraph(
        "This doesn't mean serverless is inherently less secure. In many ways, the opposite is true—AWS "
        "handles operating system patching, runtime updates, and infrastructure hardening automatically. "
        "But the application code and its dependencies remain the customer's responsibility under the "
        "shared responsibility model. Those packages imported from npm, PyPI, or Maven Central may contain "
        "vulnerabilities that put your applications at risk."
    )

    create_image_placeholder(doc, "AWS Shared Responsibility Model for Lambda: AWS responsibilities (infrastructure, runtime, OS) vs Customer responsibilities (code, dependencies, configuration)")

    # Section 2: Understanding the Landscape
    doc.add_heading('Understanding the Lambda Vulnerability Landscape', level=1)

    doc.add_paragraph(
        "To effectively secure Lambda functions, it helps to understand where vulnerabilities typically "
        "originate. Our research and customer deployments have identified three primary categories:"
    )

    # Dependency vulnerabilities
    p1 = doc.add_paragraph()
    p1_run = p1.add_run("Dependency vulnerabilities ")
    p1_run.bold = True
    p1.add_run(
        "represent the most common finding. Modern applications rely heavily on third-party packages, "
        "and those packages have their own dependencies. A typical Node.js Lambda function may include "
        "hundreds of transitive dependencies, any of which could contain known CVEs. The Log4Shell "
        "vulnerability (CVE-2021-44228) demonstrated how a single library vulnerability could impact "
        "millions of applications across the industry."
    )

    # Secrets exposure
    p2 = doc.add_paragraph()
    p2_run = p2.add_run("Secrets and sensitive data exposure ")
    p2_run.bold = True
    p2.add_run(
        "occurs more frequently than organizations expect. API keys, database credentials, and "
        "authentication tokens sometimes find their way into application code or configuration. While "
        "developers generally understand they shouldn't commit secrets to version control, the pressure "
        "of rapid deployment cycles can lead to shortcuts that create risk."
    )

    # Configuration weaknesses
    p3 = doc.add_paragraph()
    p3_run = p3.add_run("Configuration weaknesses ")
    p3_run.bold = True
    p3.add_run(
        "in the function itself—overly permissive IAM roles, missing encryption settings, or insecure "
        "environment variable handling—can amplify the impact of other vulnerabilities. A compromised "
        "function with excessive permissions poses far greater risk than one following least-privilege principles."
    )

    doc.add_paragraph(
        "AWS Inspector provides Lambda scanning capabilities that address some of these concerns, and "
        "organizations already using Inspector should continue to leverage it. However, enterprises with "
        "advanced security requirements may find certain limitations—functions encrypted with customer-managed "
        "KMS keys aren't supported, and functions not invoked within 90 days fall outside scanning scope. "
        "Organizations managing large serverless estates across multiple accounts often need additional capabilities."
    )

    # Section 3: Event-Driven Approach
    doc.add_heading('An Event-Driven Approach to Lambda Security', level=1)

    doc.add_paragraph(
        "Qualys has developed a Lambda scanning solution that integrates vulnerability detection directly "
        "into the deployment workflow. Rather than relying on scheduled scans or periodic assessments, "
        "this approach uses AWS's native event system to trigger security analysis the moment functions change."
    )

    create_image_placeholder(doc, "Architecture diagram: Lambda deployment → CloudTrail → EventBridge → Qualys Scanner Lambda → Results (S3, SNS, DynamoDB)")

    doc.add_paragraph(
        "The architecture leverages familiar AWS services. When a Lambda function is created or updated, "
        "AWS CloudTrail records the API call. Amazon EventBridge routes matching events to a scanner "
        "function that performs the security assessment. Results flow to Amazon S3 for retention, Amazon "
        "SNS for notifications, and Amazon DynamoDB for caching."
    )

    doc.add_paragraph(
        "This event-driven model offers several advantages over scheduled scanning. Coverage is immediate—"
        "there's no window between deployment and assessment. Resources are used efficiently—functions are "
        "scanned when they change, not on arbitrary schedules. And the approach scales naturally with your "
        "serverless adoption; whether you deploy ten functions per month or ten thousand, the scanner "
        "responds proportionally."
    )

    doc.add_paragraph(
        "The scanning engine itself uses Qualys's vulnerability intelligence database, the same research "
        "that powers our broader vulnerability management platform. This ensures consistency across your "
        "security program—Lambda vulnerabilities are identified, categorized, and prioritized using the "
        "same framework applied to your containers, virtual machines, and other assets."
    )

    # Section 4: Deployment Models
    doc.add_heading('Deployment Models for Different Organizational Needs', level=1)

    doc.add_paragraph(
        "Enterprise environments vary significantly in their AWS account structure and operational models. "
        "We've designed the solution to accommodate this diversity through three deployment approaches."
    )

    # Single account
    p_single = doc.add_paragraph()
    p_single_run = p_single.add_run("Single-account deployment ")
    p_single_run.bold = True
    p_single.add_run(
        "suits organizations with contained AWS footprints or those beginning their serverless security "
        "journey. All components run within a single account, providing complete coverage with minimal "
        "configuration complexity. This model works well for development teams or smaller organizations "
        "where centralized AWS management is the norm."
    )

    # Hub-spoke
    p_hub = doc.add_paragraph()
    p_hub_run = p_hub.add_run("Hub-spoke architecture ")
    p_hub_run.bold = True
    p_hub.add_run(
        "addresses the needs of organizations with centralized security operations. A scanner running "
        "in a dedicated security account reaches into member accounts using carefully scoped cross-account "
        "IAM roles. EventBridge rules in spoke accounts forward Lambda events to the central hub, creating "
        "unified visibility regardless of where functions are deployed. This model aligns well with "
        "organizations that have established security account patterns."
    )

    # StackSet
    p_stack = doc.add_paragraph()
    p_stack_run = p_stack.add_run("StackSet distribution ")
    p_stack_run.bold = True
    p_stack.add_run(
        "provides an alternative for organizations that prefer operational independence across accounts. "
        "CloudFormation StackSets deploy consistent scanner infrastructure to every account in an AWS "
        "Organization, with each account maintaining its own scanning capability. This approach reduces "
        "cross-account dependencies while ensuring consistent security coverage."
    )

    create_image_placeholder(doc, "Three deployment models: Single Account (simple), Hub-Spoke (centralized), StackSet Distribution (federated)")

    doc.add_paragraph(
        "All deployment models implement encryption by default—KMS for data at rest, enforced TLS for "
        "data in transit, Secrets Manager for credentials. IAM policies follow least-privilege principles, "
        "with permissions scoped to specific resources rather than using wildcards. These security controls "
        "are embedded in the infrastructure code, ensuring consistent implementation across deployments."
    )

    # Section 5: QFlow Integration
    doc.add_heading('From Detection to Response with QFlow Integration', level=1)

    doc.add_paragraph(
        "Identifying vulnerabilities is a necessary first step, but organizations increasingly recognize "
        "that detection alone is insufficient. The time between discovering a critical vulnerability and "
        "remediating it represents active risk exposure. Shortening this window has become a priority for "
        "security programs."
    )

    doc.add_paragraph(
        "QFlow integration enables automated response workflows that can dramatically reduce exposure time. "
        "When the scanner identifies a vulnerability exceeding defined severity thresholds, QFlow can "
        "execute predefined response actions without waiting for human intervention."
    )

    doc.add_paragraph(
        "Consider a practical scenario: a critical remote code execution vulnerability is detected in a "
        "production Lambda function at 2 AM on a Saturday. In a traditional workflow, this finding waits "
        "in a queue until Monday morning. With QFlow, automated response initiates immediately."
    )

    create_image_placeholder(doc, "QFlow automated response workflow: Vulnerability detected → Severity evaluation → Automated isolation → Notifications triggered → Ticket created")

    # Isolation techniques
    doc.add_heading('Automated Isolation Techniques', level=2)

    p_concurrency = doc.add_paragraph()
    p_concurrency_run = p_concurrency.add_run("Isolation through concurrency control ")
    p_concurrency_run.bold = True
    p_concurrency.add_run(
        "provides an immediate containment option. Setting a function's reserved concurrency to zero "
        "prevents any new invocations while preserving the function's code and configuration for "
        "investigation. This is analogous to taking a compromised server offline without destroying "
        "forensic evidence."
    )

    p_permission = doc.add_paragraph()
    p_permission_run = p_permission.add_run("Permission restriction ")
    p_permission_run.bold = True
    p_permission.add_run(
        "offers a more aggressive response when warranted. Attaching a deny-all IAM policy to the "
        "function's execution role ensures that even if the function somehow executes, it cannot access "
        "AWS services or data. The function becomes effectively sandboxed."
    )

    p_network = doc.add_paragraph()
    p_network_run = p_network.add_run("Network isolation ")
    p_network_run.bold = True
    p_network.add_run(
        "applies to VPC-attached functions. Moving the function to a quarantine security group with no "
        "ingress or egress rules severs all network connectivity, limiting potential lateral movement "
        "or data exfiltration."
    )

    doc.add_paragraph(
        "These automated responses don't replace human judgment—they buy time. Security teams can "
        "investigate findings and make informed remediation decisions without the pressure of active "
        "exploitation risk."
    )

    # Section 6: Operational Considerations
    doc.add_heading('Operational Considerations', level=1)

    doc.add_paragraph(
        "Organizations implementing Lambda security programs should consider several operational factors "
        "that affect success."
    )

    p_cache = doc.add_paragraph()
    p_cache_run = p_cache.add_run("Caching and efficiency ")
    p_cache_run.bold = True
    p_cache.add_run(
        "matter at scale. The scanner maintains a cache of previously analyzed function code hashes. "
        "When a function is deployed without code changes—perhaps due to configuration updates or "
        "permission modifications—the cached assessment is used rather than rescanning. This significantly "
        "reduces scanning volume while maintaining accurate coverage."
    )

    p_alert = doc.add_paragraph()
    p_alert_run = p_alert.add_run("Alert management ")
    p_alert_run.bold = True
    p_alert.add_run(
        "requires thoughtful configuration. Lambda-heavy environments may generate substantial scanning "
        "activity. Configuring appropriate severity thresholds for notifications, leveraging SNS filtering, "
        "and integrating with existing SIEM platforms helps security teams focus on findings that require "
        "attention rather than being overwhelmed by volume."
    )

    p_baseline = doc.add_paragraph()
    p_baseline_run = p_baseline.add_run("Baseline establishment ")
    p_baseline_run.bold = True
    p_baseline.add_run(
        "for existing environments deserves attention. Organizations with established Lambda footprints "
        "should plan for initial bulk scanning to understand their current vulnerability posture. A "
        "dedicated bulk scan function can enumerate all Lambda functions across specified accounts and "
        "regions, triggering assessments systematically while respecting rate limits."
    )

    p_integration = doc.add_paragraph()
    p_integration_run = p_integration.add_run("Integration with development workflows ")
    p_integration_run.bold = True
    p_integration.add_run(
        "extends security impact. While the solution operates independently of CI/CD pipelines, many "
        "organizations find value in connecting scan results to development processes—failing builds "
        "when critical vulnerabilities are detected, creating tickets automatically, or notifying "
        "developers directly through collaboration platforms."
    )

    create_image_placeholder(doc, "Operational dashboard: Scan success rates, average duration, vulnerabilities by severity over time, cache hit ratio")

    # Section 7: Customer Outcomes
    doc.add_heading('Customer Outcomes', level=1)

    doc.add_paragraph(
        "Organizations that have implemented this approach report meaningful improvements in their "
        "serverless security posture."
    )

    doc.add_paragraph(
        "A financial services firm managing over 500 Lambda functions across 12 AWS accounts reduced "
        "their mean time to identify Lambda vulnerabilities from their previous 30-day scan cycle to "
        "near-real-time detection. More significantly, QFlow automation reduced their mean time to "
        "contain critical findings from days to minutes."
    )

    doc.add_paragraph(
        "A healthcare organization used the solution to achieve continuous compliance monitoring for "
        "Lambda workloads, generating audit evidence automatically rather than through periodic manual "
        "assessments. Their compliance team appreciated the historical scan records stored in S3 and "
        "the function-level tagging that provides immediate visibility into security status."
    )

    doc.add_paragraph(
        "A software company integrated scanning results into their deployment pipeline, preventing "
        "vulnerable code from reaching production environments. Developers received faster feedback "
        "on security issues, and the security team shifted focus from finding vulnerabilities to "
        "preventing them."
    )

    doc.add_paragraph(
        "These outcomes reflect a broader trend: organizations are moving from periodic assessment "
        "to continuous security operations, and extending this capability to serverless workloads "
        "is a natural evolution."
    )

    # Section 8: Looking Forward
    doc.add_heading('Looking Forward', level=1)

    doc.add_paragraph(
        "Serverless computing will continue growing as organizations pursue operational efficiency and "
        "development velocity. The security programs that protect these workloads must evolve alongside them."
    )

    doc.add_paragraph(
        "The approach we've described—event-driven scanning integrated with automated response—represents "
        "one step in this evolution. Future developments may include deeper integration with runtime "
        "behavior analysis, expanded coverage for additional serverless services, and more sophisticated "
        "policy engines that consider business context alongside technical findings."
    )

    doc.add_paragraph(
        "What remains constant is the fundamental principle: security should enable rather than impede "
        "cloud adoption. By making vulnerability detection automatic and response immediate, organizations "
        "can embrace serverless architectures with confidence that their security program keeps pace."
    )

    create_image_placeholder(doc, "Vision: Unified cloud security across Lambda, containers, and traditional infrastructure")

    # Getting Started section
    add_horizontal_line(doc.add_paragraph())

    doc.add_heading('Getting Started', level=1)

    doc.add_paragraph(
        "Organizations interested in extending their vulnerability management program to Lambda workloads "
        "can begin with the resources below. The solution deploys through CloudFormation templates or "
        "Terraform modules, with typical deployment completing in under 15 minutes."
    )

    resources = doc.add_paragraph()
    resources.add_run("Documentation: ").bold = True
    resources.add_run("Comprehensive deployment guides, configuration options, and operational runbooks\n")
    resources.add_run("Source Repository: ").bold = True
    resources.add_run("Infrastructure code, Lambda function source, and example configurations\n")
    resources.add_run("Qualys Community: ").bold = True
    resources.add_run("Discussion forums, best practices, and customer examples")

    doc.add_paragraph(
        "For organizations requiring assistance with design or deployment, Qualys professional services "
        "and certified partners can help plan implementations aligned with your specific environment "
        "and requirements."
    )

    # Footer
    add_horizontal_line(doc.add_paragraph())

    # Author section
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.LEFT
    author_run = author.add_run("About the Author")
    author_run.bold = True
    author_run.font.size = Pt(12)

    doc.add_paragraph("[Author name, title, and bio placeholder]")

    # Related resources
    related = doc.add_paragraph()
    related_run = related.add_run("Related Resources")
    related_run.bold = True
    related_run.font.size = Pt(12)

    doc.add_paragraph("• Qualys Cloud Agent for Container Security")
    doc.add_paragraph("• Understanding Continuous Threat Exposure Management")
    doc.add_paragraph("• AWS Security Best Practices for Serverless Applications")

    # Save the document
    doc.save('/Users/andrew/git_base/qualys-lambda/Qualys_Lambda_Scanner_Blog.docx')
    print("Document saved: Qualys_Lambda_Scanner_Blog.docx")

if __name__ == "__main__":
    main()
